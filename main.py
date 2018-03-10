#!/usr/bin/env python

import csv
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

FILE_NAME = 'worklog.csv'
HEADER_DATE = 'Start Time'
HEADER_SPENT = 'Time Spent (s)'
DOCX_FILE_NAME = 'rozliczenie.docx'

def seconds_to_full_hours(time):
    return int(time / 60 / 60)

with open(FILE_NAME, encoding="utf8") as csvfile:
    total_time = 0  # [SECONDS]
    worklogs = {}

    reader = csv.DictReader(csvfile)
    for row in reader:
        worklog_time = int(row[HEADER_SPENT])
        total_time += worklog_time
        worklog_date = datetime.strptime(row[HEADER_DATE],
                                         '%Y-%m-%d %H:%M').date()
        worklogDateStr = worklog_date.strftime('%d.%m.%Y')

        if worklogDateStr in worklogs:
            worklogs[worklogDateStr] += worklog_time
        else:
            worklogs[worklogDateStr] = worklog_time

    for key, value in worklogs.items():
        worklogs[key] = seconds_to_full_hours(worklogs[key])

    for key, value in worklogs.items():
        print(key, value)

    total_time = seconds_to_full_hours(total_time)
    print('Total time: ' + str(total_time))

    document = Document(DOCX_FILE_NAME)
    table = document.tables[0]

    worklogs_dates_list = list(worklogs.keys())
    worklogs_times_list = list(worklogs.values())
    for index, row in enumerate(table.rows[1:-1]):
        if index >= len(worklogs_dates_list):
            break

        date_cell = row.cells[0]
        paragraph = date_cell.paragraphs[0]
        paragraph.text = str(index + 1) + '. ' + str(worklogs_dates_list[index])

        hours_cell = row.cells[1]
        paragraph = hours_cell.paragraphs[0]
        paragraph.text = str(worklogs_times_list[index])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.rows[len(table.rows) - 1]
    hours_cell = row.cells[1]
    paragraph = hours_cell.paragraphs[0]
    paragraph.text = str(total_time)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save('rozliczenie-wypelnione.docx')
