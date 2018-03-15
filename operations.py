#!/usr/bin/env python

from csv import DictReader
from collections import OrderedDict
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def seconds_to_full_hours(time):
    return int(time / 60 / 60)


def process_csv(csvfile, header_date, header_spent):
    worklogs = {}

    reader = DictReader(csvfile)
    for row in reader:
        worklog_time = int(row[header_spent])
        worklog_date = datetime.strptime(row[header_date],
                                         '%Y-%m-%d %H:%M').date()
        worklogDateStr = worklog_date.strftime('%d.%m.%Y')

        if worklogDateStr in worklogs:
            worklogs[worklogDateStr] += worklog_time
        else:
            worklogs[worklogDateStr] = worklog_time

    return OrderedDict(sorted(worklogs.items(), key=lambda t: t[0]))


def process_docx(fileName, worklogs, total_time):
    document = Document(fileName)
    table = document.tables[0]

    worklogs_dates_list = list(worklogs.keys())
    worklogs_times_list = list(worklogs.values())
    for index, row in enumerate(table.rows[1:-1]):
        if index >= len(worklogs_dates_list):
            break

        date_cell = row.cells[0]
        paragraph = date_cell.paragraphs[0]
        paragraph.text = str(index + 1) + '. ' + str(
            worklogs_dates_list[index])

        hours_cell = row.cells[1]
        paragraph = hours_cell.paragraphs[0]
        paragraph.text = str(worklogs_times_list[index])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.rows[len(table.rows) - 1]
    hours_cell = row.cells[1]
    paragraph = hours_cell.paragraphs[0]
    paragraph.text = str(total_time)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save('rozliczenie-wypelnione.docx')